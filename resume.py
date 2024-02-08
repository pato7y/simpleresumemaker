from docx import Document

def create_resume():
 # Create a new Document
    doc = Document()

    # Add heading
    doc.add_heading('CV', level=1)

 # Add personal information here
    doc.add_heading('Personal Information', level=2)
    name = input('Enter your name and surname : ')
    email = input('Enter your email: ')
    phone = input('Enter your phone number: ')
    doc.add_paragraph(f'Name: {name}')
    doc.add_paragraph(f'Email: {email}')
    doc.add_paragraph(f'Phone: {phone}')

 # Add education details here
    doc.add_heading('Education', level=2)
    education = []
    while True:
        degree = input('Enter your degree: ')
        if not degree:
            break
        major = input('Enter your major: ')
        school = input('Enter your school: ')
        year = input('Enter graduation year: ')
        gpa = input('Enter your GPA: ')
        education.append({'degree': degree, 'major': major, 'school': school, 'year': year, 'gpa': gpa})
        print("if you dont wanna add more education to here just click enter ")
    for edu in education:
        doc.add_paragraph(f'{edu["degree"]} in {edu["major"]}')
        doc.add_paragraph(f'{edu["school"]}, {edu["year"]} - GPA: {edu["gpa"]}')

# Add work experience here
    doc.add_heading('Experience', level=2)
    experience = []
    while True:
        print("if you dont wanna add position you can skip  here just click enter ")
        position = input('Enter your position: ')
        if not position:
            break
        company = input('Enter company name: ')
        start_date = input('Enter start date (YYYY-MM): ')
        end_date = input('Enter end date (YYYY-MM or "Present"): ')
        description = input('Enter job description: ')
        experience.append({'position': position, 'company': company, 'start_date': start_date, 'end_date': end_date, 'description': description})
    for exp in experience:
        doc.add_paragraph(f'{exp["position"]}')
        doc.add_paragraph(f'{exp["company"]}, {exp["start_date"]} - {exp["end_date"]}')
        doc.add_paragraph(f'{exp["description"]}')

 # Adding  skills here 
    doc.add_heading('Skills', level=2)
    skills = input('Enter your skills (comma-separated): ').split(',')
    for skill in skills:
        doc.add_paragraph(skill.strip())
 # Save the document
    doc.save('resume.docx')
    print('Resume created successfully!')

#Create the resume
create_resume()
